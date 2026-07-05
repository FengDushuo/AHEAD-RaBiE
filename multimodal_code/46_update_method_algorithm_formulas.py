from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt
from lxml import etree


DOCX_PATH = Path(
    r"D:\a_work\1-phD\project\3-lin-4\enzyme_inspired_pipeline_complete\Method-revised-multiview-addhout.docx"
)
BACKUP_PATH = DOCX_PATH.with_name("Method-revised-multiview-addhout.before_algorithm_formulas.docx")
OUT_PATH = DOCX_PATH.with_name("Method-revised-multiview-addhout-algorithm-formulas.docx")
MML2OMML_XSL = Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL")
MATH_NS = "http://www.w3.org/1998/Math/MathML"


SECTION_TITLE = "Multi-view fine-tuning and target-domain prediction of hydrogen adsorption energy"
DFT_TITLE = "Density functional theory (DFT) calculations"


def mathml(children: str) -> str:
    return f'<math xmlns="{MATH_NS}" display="block">{children}</math>'


EQUATIONS = [
    mathml(
        """
        <mrow>
          <msubsup><mi mathvariant="bold">z</mi><mi>i</mi><mrow><mo>(</mo><mi>v</mi><mo>)</mo></mrow></msubsup>
          <mo>=</mo>
          <msub><mi>f</mi><mi>v</mi></msub>
          <mo>(</mo>
          <msubsup><mi mathvariant="bold">x</mi><mi>i</mi><mrow><mo>(</mo><mi>v</mi><mo>)</mo></mrow></msubsup>
          <mo>;</mo><msub><mi>θ</mi><mi>v</mi></msub>
          <mo>)</mo>
          <mo>,</mo><mspace width="0.5em"/>
          <msub><mi mathvariant="bold">z</mi><mi>i</mi></msub>
          <mo>=</mo>
          <mi>Φ</mi>
          <mo>(</mo>
          <mo>[</mo>
          <msubsup><mi mathvariant="bold">z</mi><mi>i</mi><mrow><mo>(</mo><mn>1</mn><mo>)</mo></mrow></msubsup>
          <mo>,</mo><mo>…</mo><mo>,</mo>
          <msubsup><mi mathvariant="bold">z</mi><mi>i</mi><mrow><mo>(</mo><mi>V</mi><mo>)</mo></mrow></msubsup>
          <mo>,</mo>
          <msub><mi mathvariant="bold">p</mi><mi>i</mi></msub>
          <mo>]</mo><mo>)</mo>
        </mrow>
        """
    ),
    mathml(
        """
        <mrow>
          <msubsup><mover><mi>y</mi><mo>^</mo></mover><mi>i</mi><mrow><mo>(</mo><mn>0</mn><mo>)</mo></mrow></msubsup>
          <mo>=</mo>
          <mi>g</mi><mo>(</mo><msub><mi mathvariant="bold">z</mi><mi>i</mi></msub><mo>;</mo><msub><mi>θ</mi><mi>g</mi></msub><mo>)</mo>
          <mo>,</mo><mspace width="0.5em"/>
          <msub><mi mathvariant="script">L</mi><mtext>src</mtext></msub>
          <mo>=</mo>
          <mfrac><mn>1</mn><msub><mi>N</mi><mi>s</mi></msub></mfrac>
          <munderover>
            <mo>∑</mo>
            <mrow><mi>i</mi><mo>=</mo><mn>1</mn></mrow>
            <msub><mi>N</mi><mi>s</mi></msub>
          </munderover>
          <msub><mi>ρ</mi><mi>δ</mi></msub>
          <mo>(</mo>
          <msub><mi>y</mi><mi>i</mi></msub>
          <mo>−</mo>
          <msubsup><mover><mi>y</mi><mo>^</mo></mover><mi>i</mi><mrow><mo>(</mo><mn>0</mn><mo>)</mo></mrow></msubsup>
          <mo>)</mo>
          <mo>+</mo>
          <mi>λ</mi>
          <msubsup>
            <mrow><mo>∥</mo><mi>Θ</mi><mo>−</mo><msub><mi>Θ</mi><mn>0</mn></msub><mo>∥</mo></mrow>
            <mn>2</mn>
            <mn>2</mn>
          </msubsup>
        </mrow>
        """
    ),
    mathml(
        """
        <mrow>
          <msubsup><mover><mi>y</mi><mo>^</mo></mover><mi>j</mi><mtext>prior</mtext></msubsup>
          <mo>=</mo>
          <msubsup><mover><mi>y</mi><mo>^</mo></mover><mi>j</mi><mrow><mo>(</mo><mn>0</mn><mo>)</mo></mrow></msubsup>
          <mo>+</mo>
          <mi>A</mi>
          <mo>(</mo><msub><mi>h</mi><mi>j</mi></msub><mo>,</mo><msub><mi>d</mi><mi>j</mi></msub><mo>,</mo><msub><mi>r</mi><mi>j</mi></msub><mo>;</mo><mi>α</mi><mo>)</mo>
          <mo>+</mo>
          <msubsup><mi mathvariant="bold">q</mi><mi>j</mi><mi>T</mi></msubsup>
          <mi mathvariant="bold">β</mi>
        </mrow>
        """
    ),
    mathml(
        """
        <mrow>
          <msup><mi mathvariant="bold">γ</mi><mo>*</mo></msup>
          <mo>=</mo>
          <munder>
            <mtext>arg min</mtext>
            <mi mathvariant="bold">γ</mi>
          </munder>
          <mo>[</mo>
          <munder>
            <mo>∑</mo>
            <mrow><mi>j</mi><mo>∈</mo><mi mathvariant="script">C</mi></mrow>
          </munder>
          <msup>
            <mrow>
              <mo>(</mo>
              <msub><mi>y</mi><mi>j</mi></msub>
              <mo>−</mo>
              <msubsup><mover><mi>y</mi><mo>^</mo></mover><mi>j</mi><mtext>prior</mtext></msubsup>
              <mo>−</mo>
              <msubsup><mi mathvariant="bold">ψ</mi><mi>j</mi><mi>T</mi></msubsup>
              <mi mathvariant="bold">γ</mi>
              <mo>)</mo>
            </mrow>
            <mn>2</mn>
          </msup>
          <mo>+</mo>
          <msub><mi>λ</mi><mi>c</mi></msub>
          <msubsup><mrow><mo>∥</mo><mi mathvariant="bold">γ</mi><mo>∥</mo></mrow><mn>2</mn><mn>2</mn></msubsup>
          <mo>]</mo>
          <mo>,</mo><mspace width="0.5em"/>
          <msubsup><mover><mi>y</mi><mo>^</mo></mover><mi>j</mi><mtext>cal</mtext></msubsup>
          <mo>=</mo>
          <msubsup><mover><mi>y</mi><mo>^</mo></mover><mi>j</mi><mtext>prior</mtext></msubsup>
          <mo>+</mo>
          <msubsup><mi mathvariant="bold">ψ</mi><mi>j</mi><mi>T</mi></msubsup>
          <msup><mi mathvariant="bold">γ</mi><mo>*</mo></msup>
        </mrow>
        """
    ),
]


NEW_PARAGRAPHS = [
    (
        "text",
        "Each catalyst configuration was treated as a sample i with V complementary descriptor views. "
        "These views included host-oxide identity, dopant identity, adsorption-site and local geometric descriptors, "
        "structure-derived graph embeddings, source-domain dopant statistics and LLM-derived elemental priors. "
        "For each view v, a view-specific encoder fv mapped the input vector xiv into a latent representation ziv, "
        "and the fused multi-view state zi was obtained by a fusion operator Φ that also incorporated the prior vector pi:",
    ),
    ("eq", 0),
    (
        "text",
        "The original multi-view model was used as the initialization and was fine-tuned on the curated source-domain "
        "hydrogen-adsorption dataset. The initial prediction yhat(0) was generated by the regression head g, and the "
        "fine-tuning objective combined a robust Huber loss ρδ with an L2 regularization term that constrained the "
        "updated parameters Θ to remain close to the pretrained parameters Θ0:",
    ),
    ("eq", 1),
    (
        "text",
        "This formulation allowed the model to preserve the transferable representation learned by the multi-view "
        "backbone while adapting the regression head and selected encoder layers to the adsorption-energy scale of "
        "the present oxide systems. To reduce source-to-target domain shift for the independent CeO2/ZnO target "
        "systems, the strict multi-view prediction was further corrected by a chemistry-guided anchor. This anchor "
        "was parameterized by the host oxide hj, dopant dj, redox/binding regime rj and a vector qj of elemental "
        "knowledge priors, including reducibility, oxygen affinity and H* binding tendency:",
    ),
    ("eq", 2),
    (
        "text",
        "When a small number of target-domain DFT labels were available, they were used only as a calibration subset C. "
        "A low-capacity ridge residual model was then fitted on C and applied to the remaining held-out target samples; "
        "the held-out labels were never used during fitting. The calibrated prediction was defined as:",
    ),
    ("eq", 3),
    (
        "text",
        "Here ψj denotes the residual calibration features, which were restricted to low-dimensional chemistry and "
        "model-output terms to avoid overfitting. Model performance was reported on the held-out target samples using "
        "mean absolute error, root-mean-square error, Pearson correlation and Spearman rank correlation under "
        "material-stratified repeated splits. This protocol separates strict source-domain transfer, chemistry-prior "
        "correction and DFT-assisted few-shot calibration, and therefore provides a controlled evaluation of whether "
        "limited DFT information can correct the AddH target-domain shift without obscuring the contribution of the "
        "multi-view model."
    ),
]


def set_run_font(run, *, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    run.font.bold = bold
    r_fonts = run._element.rPr.rFonts
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "Times New Roman")


def set_text_paragraph_format(paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = None


def add_text_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    set_text_paragraph_format(p)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_equation_paragraph(doc: Document, mathml_text: str, transform):
    p = doc.add_paragraph()
    p.style = "Normal"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    omml = transform(etree.fromstring(mathml_text.encode("utf-8"))).getroot()
    p._p.append(omml)
    return p


def remove_between(section_title_paragraph, dft_title_paragraph) -> None:
    body = section_title_paragraph._p.getparent()
    current = section_title_paragraph._p.getnext()
    while current is not None and current is not dft_title_paragraph._p:
        nxt = current.getnext()
        body.remove(current)
        current = nxt


def main() -> None:
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except AttributeError:
        pass

    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)
    if not MML2OMML_XSL.exists():
        raise FileNotFoundError(f"MathML-to-OMML converter not found: {MML2OMML_XSL}")
    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    transform = etree.XSLT(etree.parse(str(MML2OMML_XSL)))
    doc = Document(str(DOCX_PATH))

    section = None
    dft = None
    for p in doc.paragraphs:
        if p.text.strip() == SECTION_TITLE:
            section = p
        elif p.text.strip() == DFT_TITLE:
            dft = p
            break

    if section is None or dft is None:
        raise RuntimeError("Could not locate the multi-view section or DFT anchor.")

    remove_between(section, dft)

    new_nodes = []
    for kind, payload in NEW_PARAGRAPHS:
        if kind == "text":
            new_nodes.append(add_text_paragraph(doc, payload)._p)
        elif kind == "eq":
            new_nodes.append(add_equation_paragraph(doc, EQUATIONS[payload], transform)._p)
        else:
            raise ValueError(kind)

    body = doc._body._element
    for node in new_nodes:
        body.remove(node)
    insert_at = body.index(dft._p)
    for offset, node in enumerate(new_nodes):
        body.insert(insert_at + offset, node)

    doc.save(str(OUT_PATH))

    # Structural audit.
    updated = Document(str(OUT_PATH))
    text = "\n".join(p.text for p in updated.paragraphs)
    print(f"updated={OUT_PATH}")
    print(f"backup={BACKUP_PATH}")
    print(f"paragraphs={len(updated.paragraphs)}")
    print(f"section_present={SECTION_TITLE in text}")
    print(f"dft_anchor_present={DFT_TITLE in text}")

    with OUT_PATH.open("rb"):
        pass


if __name__ == "__main__":
    main()
