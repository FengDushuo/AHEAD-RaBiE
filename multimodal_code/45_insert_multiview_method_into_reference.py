from __future__ import annotations

from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt


REF_DIR = Path(r"D:\a_work\1-phD\project\3-lin-4\enzyme_inspired_pipeline_complete")
OUT_NAME = "Method-revised-multiview-addhout.docx"


HEADING_TEXT = "Multi-view fine-tuning and target-domain prediction of hydrogen adsorption energy"

BODY_TEXT = (
    "To connect the literature-derived hydrogen-adsorption descriptors with catalyst-level prediction, "
    "a multi-view regression module was introduced to predict ∆EH* for doped CeO2 and ZnO surfaces. "
    "The original multi-view model was used as the backbone and fine-tuned using the curated source-domain "
    "adsorption-energy dataset. Each sample was represented by complementary views, including "
    "host oxide and dopant identity, adsorption-site and local geometric descriptors, structure-derived graph "
    "embeddings, source-domain dopant statistics, and LLM-derived elemental priors related to reducibility, "
    "oxygen affinity and H* binding tendency. View-specific representations were integrated by the multi-view "
    "learner to obtain an initial ∆EH* prediction. To reduce the source-to-target domain shift for the independent "
    "CeO2/ZnO target systems, this prediction was further corrected by a chemistry-guided target-domain anchor based "
    "on host identity, dopant class and redox/binding regime. When a small number of target-domain labels were "
    "available, only the calibration subset was used to fit a low-capacity residual correction relative to this "
    "anchor, whereas the remaining target samples were kept as held-out targets. Model performance was "
    "evaluated by material-stratified repeated splits using mean absolute error, root-mean-square error, "
    "Pearson correlation and Spearman rank correlation, thereby separating strict source-domain transfer from "
    "few-shot target-domain calibration."
)


def set_run_font(run, *, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(10.5)
    run.font.bold = bold
    r_fonts = run._element.rPr.rFonts
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:eastAsia"), "Times New Roman")


def set_paragraph_format(paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = None


def add_formatted_paragraph(doc: Document, text: str, *, bold: bool) -> object:
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold)
    return paragraph


def main() -> None:
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except AttributeError:
        pass

    candidates = sorted(
        (p for p in REF_DIR.glob("Method-*.docx") if p.name != OUT_NAME),
        key=lambda p: p.stat().st_mtime,
        reverse=True,
    )
    if not candidates:
        raise FileNotFoundError(f"No Method-*.docx found under {REF_DIR}")

    source = candidates[0]
    output = REF_DIR / OUT_NAME
    doc = Document(str(source))

    existing = "\n".join(p.text for p in doc.paragraphs)
    if HEADING_TEXT in existing:
        raise RuntimeError("The multi-view method section is already present in the selected document.")

    insertion_anchor = None
    for paragraph in doc.paragraphs:
        if "Density functional theory (DFT) calculations" in paragraph.text:
            insertion_anchor = paragraph
            break
    if insertion_anchor is None:
        raise RuntimeError("Could not find the DFT section anchor for insertion.")

    heading = add_formatted_paragraph(doc, HEADING_TEXT, bold=True)
    body = add_formatted_paragraph(doc, BODY_TEXT, bold=False)

    doc_body = doc._body._element
    doc_body.remove(heading._p)
    doc_body.remove(body._p)
    insertion_anchor._p.addprevious(heading._p)
    insertion_anchor._p.addprevious(body._p)

    doc.save(str(output))

    updated = Document(str(output))
    for idx, paragraph in enumerate(updated.paragraphs):
        if HEADING_TEXT in paragraph.text:
            context = updated.paragraphs[max(0, idx - 1) : min(len(updated.paragraphs), idx + 4)]
            print(f"source={source}")
            print(f"output={output}")
            print(f"inserted_at_paragraph={idx}")
            print("context:")
            for item in context:
                print(item.text[:240])
            break


if __name__ == "__main__":
    main()
