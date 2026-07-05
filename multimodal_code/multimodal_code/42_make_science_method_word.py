#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build a Science-style DOCX draft for the AddH-out multi-view method."""
from __future__ import annotations

from pathlib import Path

from lxml import etree
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "outputs_addh_manuscript_text"
OUT_DOCX = OUT_DIR / "addhout_multiview_methods_results_science_draft.docx"
FIG1_PNG = ROOT / "outputs_addh_top_journal_fig1" / "fig1_top_journal_multiview_addhout.png"
MML2OMML_XSL = Path(r"C:\Program Files\Microsoft Office\root\Office16\MML2OMML.XSL")


MATH_NS = "http://www.w3.org/1998/Math/MathML"
_MML_TRANSFORM = etree.XSLT(etree.parse(str(MML2OMML_XSL))) if MML2OMML_XSL.exists() else None


def mml(children: str) -> str:
    return f'<math xmlns="{MATH_NS}" display="block">{children}</math>'


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9DEE8", size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def add_paragraph(doc: Document, text: str = "", style: str | None = None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_mathml_equation(doc: Document, label: str, mathml: str, explanation: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lab = p.add_run(f"{label}   ")
    lab.font.size = Pt(9.5)
    lab.font.bold = True
    if _MML_TRANSFORM is None:
        fallback = p.add_run("[MathML-to-OMML converter unavailable]")
        fallback.font.size = Pt(10.5)
    else:
        omml = _MML_TRANSFORM(etree.fromstring(mathml.encode("utf-8"))).getroot()
        p._p.append(omml)
    if explanation:
        p2 = doc.add_paragraph(explanation)
        p2.style = "FormulaExplanation"


def add_key_table(doc: Document) -> None:
    rows = [
        ("Strict source-domain baseline", "Multi-view model fine-tuned only on source-domain labels; no AddH-out labels used."),
        ("Chemistry-guided anchor", "Target-domain correction based on host oxide, dopant identity, and chemically defined redox/binding regimes."),
        ("Few-shot calibrated model", "Residual correction learned only from the calibration subset in each AddH-out split; all remaining AddH-out samples are held out."),
        ("Main claim", "AddH-out domain shift can be reduced by combining multi-view fine-tuning, chemistry-guided anchoring, and frozen few-shot calibration."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.05)
    table.columns[1].width = Inches(4.25)
    hdr = table.rows[0].cells
    hdr[0].text = "Component"
    hdr[1].text = "Role in the manuscript"
    for c in hdr:
        set_cell_shading(c, "F2F4F7")
        set_cell_margins(c)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9.5)
    for a, b in rows:
        cells = table.add_row().cells
        cells[0].text = a
        cells[1].text = b
        for c in cells:
            set_cell_margins(c)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in c.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    r.font.size = Pt(9.3)
    set_table_borders(table)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = RGBColor(17, 17, 17)
    title.paragraph_format.space_after = Pt(4)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = RGBColor(85, 85, 85)
    subtitle.paragraph_format.space_after = Pt(12)

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 16, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 12, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)

    if "FormulaExplanation" not in styles:
        st = styles.add_style("FormulaExplanation", 1)
        st.base_style = styles["Normal"]
        st.font.size = Pt(9.3)
        st.font.color.rgb = RGBColor(85, 85, 85)
        st.paragraph_format.left_indent = Inches(0.32)
        st.paragraph_format.right_indent = Inches(0.32)
        st.paragraph_format.space_after = Pt(6)


def build_doc() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    configure_styles(doc)

    title = doc.add_paragraph(style="Title")
    title.add_run("Multi-view fine-tuning for CeO2/ZnO H adsorption-energy prediction")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Science-style manuscript draft: main-text framing, Materials and Methods, equations, and Fig. 1 caption")

    add_paragraph(
        doc,
        "This draft is written to emphasize the central methodological idea: a multi-view predictor is fine-tuned on curated addH/addH-2 source-domain data and then adapted for CeO2/ZnO AddH-out prediction using chemistry-guided anchoring and few-shot target-domain residual calibration. The description is framed at the level of the scientific model, training objective, validation protocol, and performance claim.",
    )

    if FIG1_PNG.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(FIG1_PNG), width=Inches(6.5))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run("Figure 1. Algorithmic schematic of the multi-view fine-tuning and CeO2/ZnO target-domain prediction strategy.")
        r.italic = True
        r.font.size = Pt(9.3)

    doc.add_heading("Main-text framing", level=1)
    add_paragraph(
        doc,
        "Accurate prediction of H adsorption energies across oxide hosts requires models that are sensitive to both local adsorption geometry and chemically meaningful host-dopant trends. We therefore used the original multi-view architecture as the backbone and fine-tuned it on our curated addH/addH-2 adsorption-energy data. The model integrates composition, local geometric descriptors, graph-based structural embeddings, source-domain dopant statistics, and elemental prior information into a unified representation for H adsorption-energy regression.",
    )
    add_paragraph(
        doc,
        "Direct transfer from the source-domain data to CeO2/ZnO AddH-out structures showed a pronounced target-domain shift. To address this, we combined the fine-tuned multi-view predictor with a chemically constrained target-domain anchor and a frozen few-shot calibration protocol. The anchor corrects systematic host- and dopant-dependent deviations, whereas the few-shot residual model estimates only the remaining calibration-fold error relative to this anchor. The remaining AddH-out samples were held out for evaluation in each repeated split.",
    )
    add_paragraph(
        doc,
        "In repeated material-stratified AddH-out validation, the strict source-domain model produced a held-out MAE of 1.667 +/- 0.350 eV and a Spearman correlation of 0.546 +/- 0.176, indicating that the dominant failure mode was systematic domain shift rather than purely random error. The chemistry-guided anchor reduced the held-out MAE to 0.795 +/- 0.139 eV and increased the Spearman correlation to 0.886 +/- 0.057. With few-shot target-domain calibration, the held-out MAE further decreased to 0.388 +/- 0.087 eV, with an RMSE of 0.492 +/- 0.112 eV, Pearson correlation of 0.981 +/- 0.011, and Spearman correlation of 0.967 +/- 0.025.",
    )
    add_paragraph(
        doc,
        "These results should be interpreted as evidence for target-domain calibration, not as a fully blind external-test result. The strict source-domain baseline quantifies the extrapolation gap from addH/addH-2 to AddH-out, whereas the calibrated model quantifies how much of this gap can be removed when a small, explicitly separated subset of AddH-out labels is available for target-domain adaptation.",
    )

    doc.add_heading("Materials and Methods", level=1)
    doc.add_heading("Prediction task and data partitioning", level=2)
    add_paragraph(
        doc,
        "The prediction target was the DFT H adsorption energy, denoted y. The addH and addH-2 datasets were treated as source-domain training data, whereas the CeO2/ZnO AddH-out set was treated as a target-domain validation set. Source labels from the two training collections were standardized onto a common adsorption-energy scale and associated with unified sample identifiers, host oxide labels, dopant identities, adsorption-site information, and structure-derived descriptors. Target-domain labels were used only in the explicitly defined few-shot calibration folds and were never used for strict source-domain model fitting.",
    )
    add_mathml_equation(
        doc,
        "Eq. 1",
        mml(
            """
            <mrow>
              <msub><mi>D</mi><mi>s</mi></msub>
              <mo>=</mo>
              <msubsup>
                <mrow><mo>{</mo><mo>(</mo><msub><mi>x</mi><mi>i</mi></msub><mo>,</mo><msub><mi>y</mi><mi>i</mi></msub><mo>)</mo><mo>}</mo></mrow>
                <mrow><mi>i</mi><mo>=</mo><mn>1</mn></mrow>
                <msub><mi>N</mi><mi>s</mi></msub>
              </msubsup>
              <mo>,</mo>
              <mspace width="0.7em"/>
              <msub><mi>D</mi><mi>t</mi></msub>
              <mo>=</mo>
              <msubsup>
                <mrow><mo>{</mo><mo>(</mo><msub><mi>x</mi><mi>j</mi></msub><mo>,</mo><msub><mi>y</mi><mi>j</mi></msub><mo>)</mo><mo>}</mo></mrow>
                <mrow><mi>j</mi><mo>=</mo><mn>1</mn></mrow>
                <msub><mi>N</mi><mi>t</mi></msub>
              </msubsup>
            </mrow>
            """
        ),
        "Here Ds denotes the source-domain addH/addH-2 training set and Dt denotes the CeO2/ZnO AddH-out target domain. In strict transfer, labels yj in Dt are hidden during model fitting.",
    )

    doc.add_heading("Multi-view representation", level=2)
    add_paragraph(
        doc,
        "Each sample was represented by V complementary views. These views included host and dopant descriptors, local adsorption-site geometry, graph-based structural embeddings from a pretrained atomistic representation model, source-domain dopant statistics, and elemental prior descriptors that encode chemically interpretable properties such as reducibility, oxygen affinity, and expected H-binding strength. The multi-view backbone maps each view to a latent representation and then fuses the view-specific embeddings with learned weights.",
    )
    add_mathml_equation(
        doc,
        "Eq. 2",
        mml(
            """
            <mtable>
              <mtr><mtd><mrow>
                <msubsup><mi>h</mi><mi>i</mi><mrow><mo>(</mo><mi>v</mi><mo>)</mo></mrow></msubsup>
                <mo>=</mo>
                <msub><mi>&#x03D5;</mi><mi>v</mi></msub>
                <mo>(</mo><msubsup><mi>x</mi><mi>i</mi><mrow><mo>(</mo><mi>v</mi><mo>)</mo></mrow></msubsup><mo>)</mo>
              </mrow></mtd></mtr>
              <mtr><mtd><mrow>
                <msubsup><mi>&#x03B1;</mi><mi>i</mi><mrow><mo>(</mo><mi>v</mi><mo>)</mo></mrow></msubsup>
                <mo>=</mo>
                <mfrac>
                  <mrow><mi>exp</mi><mo>(</mo><msup><mi>u</mi><mi>T</mi></msup><msubsup><mi>h</mi><mi>i</mi><mrow><mo>(</mo><mi>v</mi><mo>)</mo></mrow></msubsup><mo>)</mo></mrow>
                  <mrow><msubsup><mo>&#x2211;</mo><mrow><mi>k</mi><mo>=</mo><mn>1</mn></mrow><mi>V</mi></msubsup><mi>exp</mi><mo>(</mo><msup><mi>u</mi><mi>T</mi></msup><msubsup><mi>h</mi><mi>i</mi><mrow><mo>(</mo><mi>k</mi><mo>)</mo></mrow></msubsup><mo>)</mo></mrow>
                </mfrac>
              </mrow></mtd></mtr>
              <mtr><mtd><mrow>
                <msub><mi>z</mi><mi>i</mi></msub>
                <mo>=</mo>
                <msubsup><mo>&#x2211;</mo><mrow><mi>v</mi><mo>=</mo><mn>1</mn></mrow><mi>V</mi></msubsup>
                <msubsup><mi>&#x03B1;</mi><mi>i</mi><mrow><mo>(</mo><mi>v</mi><mo>)</mo></mrow></msubsup>
                <msubsup><mi>h</mi><mi>i</mi><mrow><mo>(</mo><mi>v</mi><mo>)</mo></mrow></msubsup>
              </mrow></mtd></mtr>
            </mtable>
            """
        ),
        "The encoder phi_v maps the v-th view to a latent vector h_i^(v). The learned attention weight alpha_i^(v) controls the contribution of each view to the fused representation z_i.",
    )
    add_mathml_equation(
        doc,
        "Eq. 3",
        mml(
            """
            <mrow>
              <msubsup><mover><mi>y</mi><mo>^</mo></mover><mi>i</mi><mi>MV</mi></msubsup>
              <mo>=</mo>
              <msub><mi>g</mi><mi>&#x03B8;</mi></msub>
              <mo>(</mo><msub><mi>z</mi><mi>i</mi></msub><mo>)</mo>
            </mrow>
            """
        ),
        "The regression head g_theta converts the fused multi-view representation into a predicted H adsorption energy.",
    )

    doc.add_heading("Fine-tuning of the multi-view backbone", level=2)
    add_paragraph(
        doc,
        "The starting multi-view backbone was fine-tuned on the curated source-domain adsorption data. To retain the useful inductive bias of the original multi-view architecture while adapting it to the present H adsorption-energy task, fine-tuning minimized a robust regression objective with weight decay toward the initialization. A Huber loss was used to reduce sensitivity to high-energy outliers while preserving quadratic behavior near small residuals.",
    )
    add_mathml_equation(
        doc,
        "Eq. 4",
        mml(
            """
            <mrow>
              <msup><mi>&#x03B8;</mi><mo>*</mo></msup>
              <mo>=</mo>
              <munder><mi>argmin</mi><mi>&#x03B8;</mi></munder>
              <mrow><mo>[</mo>
                <mfrac><mn>1</mn><msub><mi>N</mi><mi>s</mi></msub></mfrac>
                <msub><mo>&#x2211;</mo><mrow><mi>i</mi><mo>&#x2208;</mo><msub><mi>D</mi><mi>s</mi></msub></mrow></msub>
                <msub><mi>&#x03C1;</mi><mi>&#x03B4;</mi></msub>
                <mo>(</mo><msub><mi>y</mi><mi>i</mi></msub><mo>-</mo><msub><mi>f</mi><mi>&#x03B8;</mi></msub><mo>(</mo><msub><mi>x</mi><mi>i</mi></msub><mo>)</mo><mo>)</mo>
                <mo>+</mo><mi>&#x03BB;</mi>
                <msubsup><mrow><mo>&#x2225;</mo><mi>&#x03B8;</mi><mo>-</mo><msub><mi>&#x03B8;</mi><mn>0</mn></msub><mo>&#x2225;</mo></mrow><mn>2</mn><mn>2</mn></msubsup>
              <mo>]</mo></mrow>
            </mrow>
            """
        ),
        "theta0 denotes the initialization of the multi-view backbone, theta* denotes the fine-tuned parameters, and f_theta is the full multi-view predictor.",
    )

    doc.add_heading("Chemistry-guided target-domain anchor", level=2)
    add_paragraph(
        doc,
        "Because CeO2/ZnO AddH-out structures occupy a shifted chemical domain relative to the source training distribution, the fine-tuned source-domain prediction was converted into a target-domain anchor using chemically constrained host- and dopant-dependent corrections. These corrections encode oxide-chemistry trends, including dopant redox activity, charge-compensation tendency, host reducibility, and expected changes in H-binding strength. The correction layer was fixed before held-out evaluation and was used as an anchor rather than as a replacement for the multi-view model.",
    )
    add_mathml_equation(
        doc,
        "Eq. 5",
        mml(
            """
            <mrow>
              <msub><mi>a</mi><mi>j</mi></msub>
              <mo>=</mo>
              <msubsup><mover><mi>y</mi><mo>^</mo></mover><mi>j</mi><mi>MV</mi></msubsup>
              <mo>+</mo>
              <mi>b</mi><mo>(</mo><msub><mi>m</mi><mi>j</mi></msub><mo>,</mo><msub><mi>d</mi><mi>j</mi></msub><mo>,</mo><msub><mi>c</mi><mi>j</mi></msub><mo>)</mo>
            </mrow>
            """
        ),
        "For target sample j, a_j is the chemistry-guided anchor, m_j is the host oxide, d_j is the dopant identity, c_j is the assigned chemical regime, and b(.) is the frozen chemistry-guided correction.",
    )

    doc.add_heading("Few-shot residual calibration for CeO2/ZnO", level=2)
    add_paragraph(
        doc,
        "For target-domain calibration, each repeated AddH-out split was divided into a calibration subset C and a held-out test subset T, stratified by host oxide. The calibration subset was used to fit a low-capacity residual model relative to the chemistry-guided anchor. This design restricts target-domain learning to a residual correction and reduces the risk that the few-shot model overfits the entire target-domain label distribution.",
    )
    add_mathml_equation(
        doc,
        "Eq. 6",
        mml(
            """
            <mrow>
              <msub><mover><mi>y</mi><mo>~</mo></mover><mi>j</mi></msub>
              <mo>=</mo>
              <msub><mi>a</mi><mi>j</mi></msub>
              <mo>+</mo>
              <msub><mi>r</mi><mi>&#x03B7;</mi></msub>
              <mo>(</mo><msub><mi>q</mi><mi>j</mi></msub><mo>)</mo>
            </mrow>
            """
        ),
        "The vector q_j contains calibration features such as host, dopant class, chemical regime, source-domain prior statistics, and anchor uncertainty summaries.",
    )
    add_mathml_equation(
        doc,
        "Eq. 7",
        mml(
            """
            <mrow>
              <msup><mi>&#x03B7;</mi><mo>*</mo></msup>
              <mo>=</mo>
              <munder><mi>argmin</mi><mi>&#x03B7;</mi></munder>
              <mrow><mo>[</mo>
              <msub><mo>&#x2211;</mo><mrow><mi>j</mi><mo>&#x2208;</mo><mi>C</mi></mrow></msub>
              <msup>
                <mrow><mo>(</mo><msub><mi>y</mi><mi>j</mi></msub><mo>-</mo><msub><mi>a</mi><mi>j</mi></msub><mo>-</mo><msub><mi>r</mi><mi>&#x03B7;</mi></msub><mo>(</mo><msub><mi>q</mi><mi>j</mi></msub><mo>)</mo><mo>)</mo></mrow>
                <mn>2</mn>
              </msup>
              <mo>+</mo><mi>&#x03B3;</mi><msubsup><mrow><mo>&#x2225;</mo><mi>&#x03B7;</mi><mo>&#x2225;</mo></mrow><mn>2</mn><mn>2</mn></msubsup>
              <mo>]</mo></mrow>
            </mrow>
            """
        ),
        "Only samples in C are used to estimate eta*. The final prediction ytilde_j is then evaluated on held-out samples j in T.",
    )

    doc.add_heading("Validation protocol and performance metrics", level=2)
    add_paragraph(
        doc,
        "The principal evaluation was repeated material-stratified held-out validation on AddH-out. For each split, calibration labels were used only for fitting the residual calibration model, and metrics were computed only on held-out target-domain samples. Calibration-fraction sensitivity was assessed by varying the fraction of AddH-out labels assigned to C. A leave-one-dopant-out analysis was additionally used to test whether the calibration strategy remained robust when target-domain dopant identities were systematically withheld.",
    )
    add_mathml_equation(
        doc,
        "Eq. 8",
        mml(
            """
            <mrow>
              <mi>MAE</mi><mo>=</mo><mfrac><mn>1</mn><mi>n</mi></mfrac>
              <msubsup><mo>&#x2211;</mo><mrow><mi>i</mi><mo>=</mo><mn>1</mn></mrow><mi>n</mi></msubsup>
              <mo>|</mo><msub><mi>y</mi><mi>i</mi></msub><mo>-</mo><msub><mover><mi>y</mi><mo>^</mo></mover><mi>i</mi></msub><mo>|</mo>
              <mo>,</mo><mspace width="0.7em"/>
              <mi>RMSE</mi><mo>=</mo>
              <msqrt><mrow><mfrac><mn>1</mn><mi>n</mi></mfrac>
              <msubsup><mo>&#x2211;</mo><mrow><mi>i</mi><mo>=</mo><mn>1</mn></mrow><mi>n</mi></msubsup>
              <msup><mrow><mo>(</mo><msub><mi>y</mi><mi>i</mi></msub><mo>-</mo><msub><mover><mi>y</mi><mo>^</mo></mover><mi>i</mi></msub><mo>)</mo></mrow><mn>2</mn></msup>
              </mrow></msqrt>
            </mrow>
            """
        ),
        None,
    )
    add_mathml_equation(
        doc,
        "Eq. 9",
        mml(
            """
            <mrow>
              <mi>Pearson</mi><mo>=</mo><mi>corr</mi><mo>(</mo><mi>y</mi><mo>,</mo><mover><mi>y</mi><mo>^</mo></mover><mo>)</mo>
              <mo>,</mo><mspace width="0.7em"/>
              <mi>Spearman</mi><mo>=</mo><mi>corr</mi><mo>(</mo><mi>rank</mi><mo>(</mo><mi>y</mi><mo>)</mo><mo>,</mo><mi>rank</mi><mo>(</mo><mover><mi>y</mi><mo>^</mo></mover><mo>)</mo><mo>)</mo>
            </mrow>
            """
        ),
        "MAE and RMSE measure absolute energy accuracy, whereas Spearman correlation quantifies whether the predicted AddH-out trend and ranking are recovered.",
    )

    doc.add_heading("Leakage control and claims", level=2)
    add_paragraph(
        doc,
        "The strict source-domain baseline was trained without AddH-out labels and provides the relevant blind-transfer reference. The few-shot calibrated model uses AddH-out labels only within each calibration fold, and its reported performance is based on the remaining held-out samples. Therefore, the few-shot result is a target-domain calibration result rather than a fully blind external-test result. Post-hoc upper-bound analyses may be reported as supplementary ablations but should not be used as the primary performance claim.",
    )
    add_key_table(doc)

    doc.add_heading("Suggested Figure 1 caption", level=1)
    add_paragraph(
        doc,
        "Figure 1. Multi-view fine-tuning and target-domain calibration strategy for CeO2/ZnO H adsorption-energy prediction. The original multi-view backbone is fine-tuned on curated addH/addH-2 source-domain adsorption-energy labels. Each sample is represented by complementary views, including host/dopant descriptors, local geometry, graph embeddings, source-domain statistics, and elemental-prior descriptors. The fine-tuned multi-view prediction is converted into a target-domain anchor using fixed chemistry-guided host/dopant corrections. For AddH-out validation, a low-capacity residual model is fitted only on the calibration subset of each material-stratified split and evaluated on the remaining held-out CeO2/ZnO samples. This protocol separates strict source-domain transfer, chemistry-guided target anchoring, and few-shot target-domain calibration.",
    )

    doc.add_heading("Recommended manuscript wording constraints", level=1)
    add_paragraph(
        doc,
        "Use: 'few-shot AddH-out target-domain calibration with repeated held-out validation.'",
        bold_prefix="Use:",
    )
    add_paragraph(
        doc,
        "Avoid: 'fully blind external AddH-out prediction' for the few-shot calibrated result.",
        bold_prefix="Avoid:",
    )

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
